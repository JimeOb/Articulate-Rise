"""
Módulo para extraer contenido de archivo Word (.docx)
Extrae estructura de unidades, temas y contenido educativo
"""

from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from loguru import logger


class WordContentExtractor:
    """Extrae contenido educativo desde archivo Word"""

    def __init__(self, docx_path: str):
        self.docx_path = Path(docx_path)
        self.doc: Optional[Document] = None
        self.content = {
            "name": "",
            "description": "",
            "units": []
        }

        if self.docx_path.exists():
            self.load_document()
        else:
            logger.warning(f"Archivo {docx_path} no encontrado")

    def load_document(self) -> bool:
        """Cargar documento Word"""
        try:
            logger.info(f"Cargando documento: {self.docx_path}")
            self.doc = Document(self.docx_path)
            logger.info(f"Documento cargado con {len(self.doc.paragraphs)} párrafos")
            return True
        except Exception as e:
            logger.error(f"Error cargando documento: {e}")
            return False

    def extract_all_content(self) -> Dict:
        """Extraer todo el contenido del documento"""
        if not self.doc:
            logger.error("Documento no cargado")
            return self.content

        try:
            # Obtener título del documento (primer párrafo)
            if self.doc.paragraphs:
                self.content["name"] = self.doc.paragraphs[0].text.strip()
                logger.info(f"Título del curso: {self.content['name']}")

            # Extraer estructura de unidades y temas
            self._parse_document_structure()

            logger.info(f"Contenido extraído: {len(self.content['units'])} unidades")
            return self.content

        except Exception as e:
            logger.error(f"Error extrayendo contenido: {e}")
            return self.content

    def _parse_document_structure(self) -> None:
        """Analizar la estructura del documento (unidades y temas)"""

        current_unit = None
        current_theme = None
        current_section = None

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detectar nivel de encabezado por formato o patrones de texto
            # Patrón: "UNIDAD 1" o "Unidad 1"
            if any(marker in text.upper() for marker in ["UNIDAD", "UNIT"]):
                # Nueva unidad
                unit_number = self._extract_number(text)
                current_unit = {
                    "number": unit_number,
                    "title": text,
                    "themes": []
                }
                self.content["units"].append(current_unit)
                current_theme = None
                logger.debug(f"Unidad encontrada: {text}")

            # Patrón: "Tema 1.1" o "Theme 1.1"
            elif any(marker in text.upper() for marker in ["TEMA", "THEME"]) and current_unit:
                # Nuevo tema/lección
                theme_number = self._extract_number(text)
                current_theme = {
                    "code": f"{current_unit['number']}.{theme_number}",
                    "title": text,
                    "content_blocks": {
                        "narrative": "",
                        "concepts": "",
                        "academic": "",
                        "video": "",
                        "interaction": ""
                    }
                }
                current_unit["themes"].append(current_theme)
                current_section = None
                logger.debug(f"Tema encontrado: {text}")

            # Detectar secciones de contenido
            elif current_theme and text.upper() in ["NARRATIVA", "NARRATIVE", "INTRODUCCIÓN", "INTRODUCTION"]:
                current_section = "narrative"

            elif current_theme and text.upper() in ["CONCEPTOS", "CONCEPTS", "KEY CONCEPTS"]:
                current_section = "concepts"

            elif current_theme and text.upper() in ["CONTENIDO ACADÉMICO", "ACADEMIC CONTENT"]:
                current_section = "academic"

            elif current_theme and text.upper() in ["VIDEO", "DESCRIPCIÓN DE VIDEO", "VIDEO DESCRIPTION"]:
                current_section = "video"

            elif current_theme and text.upper() in ["INTERACCIÓN", "INTERACTION", "ACTIVIDAD", "ACTIVITY"]:
                current_section = "interaction"

            # Acumular contenido en la sección actual
            elif current_theme and current_section:
                current_content = current_theme["content_blocks"][current_section]
                current_theme["content_blocks"][current_section] = (current_content + " " + text).strip()

    def _extract_number(self, text: str) -> int:
        """Extraer número de texto como 'Unidad 1' o 'Tema 1.1'"""
        import re

        # Buscar patrón de número
        match = re.search(r'(\d+(?:\.\d+)?)', text)
        if match:
            num_str = match.group(1)
            # Retornar solo el primer número
            return int(num_str.split('.')[0])
        return 1

    def get_course_structure(self) -> Dict:
        """Obtener estructura formateada para automatización"""

        structure = {
            "name": self.content["name"],
            "description": self.content.get("description", ""),
            "units": []
        }

        for unit in self.content["units"]:
            unit_data = {
                "number": unit["number"],
                "title": unit["title"],
                "lessons": []
            }

            for theme in unit["themes"]:
                lesson_data = {
                    "number": theme["code"],
                    "title": theme["title"],
                    "content": self._build_content_blocks(theme)
                }
                unit_data["lessons"].append(lesson_data)

            structure["units"].append(unit_data)

        return structure

    def _build_content_blocks(self, theme: Dict) -> List[Dict]:
        """Construir bloques de contenido para una lección"""

        blocks = []
        content_blocks = theme["content_blocks"]

        if content_blocks.get("narrative"):
            blocks.append({
                "type": "Text",
                "title": "Narrativa/Introducción",
                "content": content_blocks["narrative"]
            })

        if content_blocks.get("concepts"):
            blocks.append({
                "type": "Text",
                "title": "Conceptos Clave",
                "content": content_blocks["concepts"]
            })

        if content_blocks.get("academic"):
            blocks.append({
                "type": "Text",
                "title": "Contenido Académico",
                "content": content_blocks["academic"]
            })

        if content_blocks.get("video"):
            blocks.append({
                "type": "Text",
                "title": "Descripción de Video",
                "content": content_blocks["video"]
            })

        if content_blocks.get("interaction"):
            blocks.append({
                "type": "Text",
                "title": "Interactividad/Actividad",
                "content": content_blocks["interaction"]
            })

        return blocks

    def export_json(self, output_path: str = "extracted_course_content.json") -> bool:
        """Exportar contenido extraído a JSON"""
        import json

        try:
            structure = self.get_course_structure()
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(structure, f, ensure_ascii=False, indent=2)

            logger.info(f"Contenido exportado a: {output_path}")
            return True

        except Exception as e:
            logger.error(f"Error exportando JSON: {e}")
            return False

    def get_plain_text(self) -> str:
        """Obtener todo el texto del documento como string"""
        if not self.doc:
            return ""

        return "\n".join(para.text for para in self.doc.paragraphs)


if __name__ == "__main__":
    # Uso de ejemplo
    extractor = WordContentExtractor("course_content.docx")
    content = extractor.extract_all_content()

    logger.info(f"Unidades: {len(content['units'])}")
    for unit in content['units']:
        logger.info(f"  {unit['title']}: {len(unit.get('themes', []))} temas")

    # Exportar a JSON
    extractor.export_json()

    # Obtener estructura formateada
    structure = extractor.get_course_structure()
    logger.info(f"Estructura lista para automatización")
